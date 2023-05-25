from docx import Document

def split_document(input_file, chars_per_file):
    document = Document(input_file)
    total_chars = 0
    file_count = 1
    current_file_chars = 0
    output_doc = Document()

    for paragraph in document.paragraphs:
        text = paragraph.text
        total_chars += len(text)

        if current_file_chars + len(text) <= chars_per_file:
            output_doc.add_paragraph().add_run(text)
            current_file_chars += len(text)
        else:
            remaining_chars = chars_per_file - current_file_chars
            output_doc.add_paragraph().add_run(text[:remaining_chars])
            output_file = f"output_{file_count}.docx"
            output_doc.save(output_file)

            output_doc = Document()
            output_doc.add_paragraph().add_run(text[remaining_chars:])
            current_file_chars = len(text) - remaining_chars
            file_count += 1

    if current_file_chars > 0:
        output_file = f"output_{file_count}.docx"
        output_doc.save(output_file)

    print(f"Split into {file_count} files.")

# Usage example
input_file = "Report.docx"  # Replace with your input file path
chars_per_file = 1000

split_document(input_file, chars_per_file)
